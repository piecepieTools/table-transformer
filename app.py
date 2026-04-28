"""
Table Transformer — Streamlit version
Steps: Upload → Select → Configure numeric → Verify → Download
"""

import streamlit as st
import pandas as pd
import io, os
from collections import Counter
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ──────────────────────────────────────────────────────────────
#  Page config
# ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Table Transformer",
    page_icon="📊",
    layout="centered",
)

# ──────────────────────────────────────────────────────────────
#  Custom CSS
# ──────────────────────────────────────────────────────────────
st.markdown("""
<style>
  html, body, [class*="css"] { font-family: 'Segoe UI', sans-serif; }
  #MainMenu, footer, header { visibility: hidden; }

  /* Constrain width for iframe / Lovable embed */
  .block-container {
    max-width: 780px !important;
    padding-left: 2rem !important;
    padding-right: 2rem !important;
    margin-left: auto !important;
    margin-right: auto !important;
  }

  .step-bar { display:flex; gap:0; margin-bottom:1.5rem; }
  .step-pill {
    flex:1; text-align:center; padding:7px 4px; font-size:12px;
    background:#F2F2F2; color:#AAAAAA; border:1px solid #E0E0E0;
  }
  .step-pill.active { background:#222; color:#FFF; font-weight:600; }
  .step-pill.done   { background:#555; color:#FFF; }

  .section-title { font-size:22px; font-weight:700; color:#111; margin-bottom:4px; }
  .section-sub   { font-size:14px; color:#777; margin-bottom:1.5rem; }

  .stats-strip {
    display:flex; background:#EFEFEF; border-radius:4px;
    padding:12px 0; margin:12px 0 20px; text-align:center;
  }
  .stat-item { flex:1; }
  .stat-val  { font-size:18px; font-weight:700; color:#111; }
  .stat-lbl  { font-size:11px; color:#888; margin-top:2px; }

  .success-box {
    background:#FFF; border:2px solid #222; border-radius:8px;
    padding:48px; text-align:center; margin-top:2rem;
  }
  .success-check { font-size:56px; }
  .success-title { font-size:24px; font-weight:700; color:#111; margin:12px 0 4px; }
  .success-sub   { font-size:14px; color:#888; }

  div.stButton > button {
    background:#222 !important; color:#FFF !important;
    border:none !important; border-radius:4px !important;
    padding:8px 20px !important; font-weight:600 !important;
  }
  div.stButton > button:hover { background:#444 !important; }

  div.stDownloadButton > button {
    background:#222 !important; color:#FFF !important;
    border:none !important; border-radius:4px !important;
    padding:10px 24px !important; font-weight:600 !important;
    font-size:15px !important;
  }
</style>
""", unsafe_allow_html=True)

# ──────────────────────────────────────────────────────────────
#  Constants
# ──────────────────────────────────────────────────────────────
CAT_THRESHOLD = 10
ARIAL = "Arial"
OSYM = {"<": "<",  "<=": "≤", ">": ">",  ">=": "≥"}
CSYM = {"<": "≥",  "<=": ">", ">": "≤",  ">=": "<"}

STEPS = ["Upload", "Select", "Configure", "Verify", "Download"]

# Ordinal suffixes helper
def ordinal(n):
    """Return e.g. 1 → '1st', 2 → '2nd', 3 → '3rd', 4 → '4th' …"""
    if 11 <= (n % 100) <= 13:
        return f"{n}th"
    return f"{n}{['th','st','nd','rd','th','th','th','th','th','th'][n % 10]}"

# ──────────────────────────────────────────────────────────────
#  Session state
# ──────────────────────────────────────────────────────────────
def _init():
    defaults = {
        "step":       1,
        "df":         None,
        "col_types":  {},
        "selected":   [],
        # configs[col]: type/label for cat; type/label/modes for numeric
        "configs":    {},
        # entries[col] = [ {"key": raw_val, "label": display_str}, ... ]
        # Used for categorical, multi_value, AND ordered_categorical.
        "entries":    {},
        # For ordered_categorical: queue of cols needing configuration
        "ord_queue":  [],
        "cur_ord":    None,
        "ord_total":  0,
        "num_queue":  [],
        "cur_num":    None,
        "num_total":  0,
        "word_bytes": None,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

_init()
S = st.session_state

# ──────────────────────────────────────────────────────────────
#  Type detection
# ──────────────────────────────────────────────────────────────
def _is_multi_value(series):
    non_null = series.dropna().astype(str)
    if non_null.empty:
        return False
    return non_null.str.contains(",").mean() >= 0.1

def _is_ordered_categorical(series):
    """
    Heuristic: integer-typed column with ≤ CAT_THRESHOLD unique values,
    OR a string column whose values are all purely numeric (e.g. '1','2','3').
    The idea is columns like treatment line stored as 1,2,3,4,5.
    """
    if pd.api.types.is_integer_dtype(series):
        return 2 <= series.nunique(dropna=True) <= CAT_THRESHOLD
    # String column where every non-null value looks like an integer
    non_null = series.dropna().astype(str).str.strip()
    if non_null.empty:
        return False
    return (non_null.str.fullmatch(r'\d+').all() and
            2 <= non_null.nunique() <= CAT_THRESHOLD)

def detect_type(series):
    if not pd.api.types.is_numeric_dtype(series):
        if _is_multi_value(series):
            return "multi_value"
        return "categorical"
    if series.nunique(dropna=True) <= CAT_THRESHOLD:
        if _is_ordered_categorical(series):
            return "ordered_categorical"
        return "categorical"
    return "numeric"

# ──────────────────────────────────────────────────────────────
#  Multi-value helpers
# ──────────────────────────────────────────────────────────────
def split_multi(series):
    terms = []
    for val in series.dropna().astype(str):
        for t in val.split(","):
            t = t.strip()
            if t:
                terms.append(t.lower())
    return terms

def multi_value_counts(series):
    return Counter(split_multi(series))

# ──────────────────────────────────────────────────────────────
#  Build initial entries list for a column
#  entries = [ {"key": raw_key, "label": display_label}, ... ]
# ──────────────────────────────────────────────────────────────
def build_entries(col, df, dtype):
    if dtype == "categorical":
        vals = list(df[col].value_counts(dropna=True).index)
        return [{"key": str(v), "label": str(v)} for v in vals]
    elif dtype == "ordered_categorical":
        # Sort numerically so 1 < 2 < ... < N
        vals = sorted(df[col].dropna().unique(), key=lambda x: float(x))
        entries = []
        for v in vals:
            n = int(float(v))
            entries.append({"key": str(v), "label": f"{ordinal(n)} line"})
        return entries
    else:  # multi_value
        counts = multi_value_counts(df[col])
        vals = sorted(counts.keys(), key=lambda k: -counts[k])
        return [{"key": v, "label": v.title()} for v in vals]

# ──────────────────────────────────────────────────────────────
#  Formatting
# ──────────────────────────────────────────────────────────────
def fmt_pct(pct):
    return f"{int(pct)}%" if pct == int(pct) else f"{pct:.1f}%"

def fmt_thresh(v):
    """Drop .0 for round numbers, keep decimals otherwise."""
    return f"{v:g}"

def go(step):
    S.step = step
    st.rerun()

# ──────────────────────────────────────────────────────────────
#  Step pill bar
# ──────────────────────────────────────────────────────────────
def render_steps():
    pills = ""
    for i, name in enumerate(STEPS, 1):
        if i < S.step:    cls = "step-pill done"
        elif i == S.step: cls = "step-pill active"
        else:             cls = "step-pill"
        pills += f'<div class="{cls}">{i}. {name}</div>'
    st.markdown(f'<div class="step-bar">{pills}</div>', unsafe_allow_html=True)

# ──────────────────────────────────────────────────────────────
#  Word export
# ──────────────────────────────────────────────────────────────
def _set_arial(run):
    rPr = run._r.get_or_add_rPr()
    rf = OxmlElement("w:rFonts")
    for k in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(k), ARIAL)
    rPr.insert(0, rf)

def _wcell(cell, text, bold=False, indent=False, center=False):
    para = cell.paragraphs[0]; para.clear()
    if center: para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent: para.paragraph_format.left_indent = Pt(16)
    run = para.add_run(text)
    run.font.name = ARIAL; run.font.size = Pt(10); run.bold = bold
    _set_arial(run)

def _clr(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for s in tcPr.findall(qn("w:shd")): tcPr.remove(s)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "auto")
    tcPr.append(shd)

def _rh(row, t=400):
    trPr = row._tr.get_or_add_trPr()
    rh = OxmlElement("w:trHeight"); rh.set(qn("w:val"), str(t)); trPr.append(rh)

def _hborder(row):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcB = tcPr.find(qn("w:tcBorders"))
        if tcB is None: tcB = OxmlElement("w:tcBorders"); tcPr.append(tcB)
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), "666666")
        tcB.append(b)

def _drow(tbl, c1, c2="", c3="", bold1=False, ind1=False):
    row = tbl.add_row(); _rh(row)
    for c in row.cells: _clr(c)
    _wcell(row.cells[0], c1, bold=bold1, indent=ind1)
    _wcell(row.cells[1], c2, center=True)
    _wcell(row.cells[2], c3, center=True)

def build_word_bytes(col_order, configs, entries, df):
    tpl = os.path.join(os.path.dirname(os.path.abspath(__file__)), "template.docx")
    if os.path.exists(tpl):
        doc = Document(tpl)
        for t in list(doc.tables):     t._element.getparent().remove(t._element)
        for p in list(doc.paragraphs): p._element.getparent().remove(p._element)
    else:
        doc = Document()

    tbl = doc.add_table(rows=1, cols=3)
    try: tbl.style = "List Table 1 Light"
    except: pass

    ws = [4536, 900, 1134]
    grid = OxmlElement("w:tblGrid")
    for w in ws:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
    tblPr = tbl._tbl.find(qn("w:tblPr")); tblPr.addnext(grid)
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None: tcW = OxmlElement("w:tcW"); tcPr.insert(0, tcW)
            tcW.set(qn("w:w"), str(ws[i])); tcW.set(qn("w:type"), "dxa")

    hdr = tbl.rows[0]; _rh(hdr, 500)
    for c in hdr.cells: _clr(c)
    _wcell(hdr.cells[0], "Variable", bold=True)
    _wcell(hdr.cells[1], "N",        bold=True, center=True)
    _wcell(hdr.cells[2], "%",        bold=True, center=True)
    _hborder(hdr)

    OPS_fn = {
        "<":  lambda s, v: s < v,
        "<=": lambda s, v: s <= v,
        ">":  lambda s, v: s > v,
        ">=": lambda s, v: s >= v,
    }

    for col in col_order:
        cfg    = configs[col]
        series = df[col]
        total  = len(series.dropna())
        dtype  = cfg["type"]

        if dtype in ("categorical", "ordered_categorical"):
            _drow(tbl, cfg["label"], "", "", bold1=True)
            for entry in entries.get(col, []):
                cnt = (series.astype(str) == str(entry["key"])).sum()
                pct = cnt / total * 100 if total else 0
                _drow(tbl, entry["label"], str(cnt), fmt_pct(pct), ind1=True)
            n_miss = series.isna().sum()
            if n_miss:
                _drow(tbl, "Missing", str(n_miss), fmt_pct(n_miss / len(series) * 100), ind1=True)

        elif dtype == "multi_value":
            _drow(tbl, cfg["label"], str(total), "", bold1=True)
            counts = multi_value_counts(series)
            for entry in entries.get(col, []):
                cnt = counts.get(entry["key"], 0)
                pct = cnt / total * 100 if total else 0
                _drow(tbl, entry["label"], str(cnt), fmt_pct(pct), ind1=True)
            n_miss = series.isna().sum()
            if n_miss:
                _drow(tbl, "Missing", str(n_miss), fmt_pct(n_miss / len(series) * 100), ind1=True)

        else:  # numeric
            for mc in cfg["modes"]:
                mode  = mc["mode"]
                label = mc["label"]
                s     = series.dropna()
                if mode == "mean_sd":
                    _drow(tbl, label, f"{s.mean():.2f} ± {s.std():.2f}", "", bold1=True)
                elif mode == "median":
                    q1, q3 = s.quantile(0.25), s.quantile(0.75)
                    _drow(tbl, label, f"{s.median():.2f} [{q1:.2f}–{q3:.2f}]", "", bold1=True)
                elif mode == "threshold":
                    op = mc["op"]; v = mc["threshold"]
                    _drow(tbl, label, "", "", bold1=True)
                    mask1 = OPS_fn[op](series, v)
                    cnt1  = mask1.sum(); pct1 = cnt1 / total * 100 if total else 0
                    mask2 = ~mask1 & series.notna()
                    cnt2  = mask2.sum(); pct2 = cnt2 / total * 100 if total else 0
                    _drow(tbl, mc["label_true"],  str(cnt1), fmt_pct(pct1), ind1=True)
                    _drow(tbl, mc["label_false"], str(cnt2), fmt_pct(pct2), ind1=True)
                    n_miss = series.isna().sum()
                    if n_miss:
                        _drow(tbl, "Missing", str(n_miss),
                              fmt_pct(n_miss / len(series) * 100), ind1=True)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ──────────────────────────────────────────────────────────────
#  STEP 1 — Upload
# ──────────────────────────────────────────────────────────────
def step_upload():
    st.markdown('<p class="section-title">Table Transformer</p>', unsafe_allow_html=True)
    st.markdown('<p class="section-sub">Build a formatted Word summary table from your Excel data.</p>',
                unsafe_allow_html=True)

    uploaded = st.file_uploader(
        "Upload your Excel file",
        type=["xlsx", "xls", "xlsm"],
        help="Supports .xlsx, .xls, .xlsm",
        label_visibility="collapsed",
    )

    if uploaded:
        try:
            df = pd.read_excel(uploaded)
        except Exception as e:
            st.error(f"Could not read file: {e}"); return
        if df.empty:
            st.warning("The file appears to be empty."); return

        S.df        = df
        S.col_types = {c: detect_type(df[c]) for c in df.columns}
        go(2)

# ──────────────────────────────────────────────────────────────
#  STEP 2 — Select columns + preview
# ──────────────────────────────────────────────────────────────
def step_select():
    df = S.df
    st.markdown('<p class="section-title">Select columns</p>', unsafe_allow_html=True)
    st.markdown('<p class="section-sub">Choose which columns to include in the summary table.</p>',
                unsafe_allow_html=True)

    selected = []
    cols_per_row = 3
    all_cols = list(df.columns)

    for row_start in range(0, len(all_cols), cols_per_row):
        row_cols = all_cols[row_start:row_start + cols_per_row]
        ui_cols  = st.columns(cols_per_row)
        for ui_col, col in zip(ui_cols, row_cols):
            with ui_col:
                dtype  = S.col_types[col]
                unique = df[col].nunique(dropna=True)
                badge_label = {"multi_value": "multi-value", "ordered_categorical": "ordered"}.get(dtype, dtype)
                checked = st.checkbox(
                    f"**{col}**", value=True, key=f"sel_{col}",
                    help=f"{badge_label} · {unique} unique values"
                )
                st.caption(f"`{badge_label}` · {unique} unique values")
                if checked:
                    selected.append(col)

    st.markdown("---")
    st.markdown("**Data preview** — first 8 rows")
    st.dataframe(df.head(8), use_container_width=True, height=220)
    st.caption(f"Showing first {min(8, len(df))} of {len(df)} rows · {len(df.columns)} columns")

    st.markdown("---")
    col_l, col_r = st.columns([1, 5])
    with col_l:
        if st.button("← Back"):
            go(1)
    with col_r:
        if st.button("Continue →", type="primary"):
            if not selected:
                st.error("Please select at least one column."); return
            S.selected = selected

            # Init entries + configs for cat/multi_value/ordered_categorical.
            # Only build if not already present so Back → re-Continue keeps edits.
            for col in selected:
                dtype = S.col_types[col]
                if dtype in ("categorical", "multi_value", "ordered_categorical"):
                    if col not in S.entries:
                        S.entries[col] = build_entries(col, df, dtype)
                    if col not in S.configs:
                        S.configs[col] = {"type": dtype, "label": col}

            # Queue ordered_categorical columns (need label configuration)
            S.ord_queue = [c for c in selected if S.col_types[c] == "ordered_categorical"]
            S.ord_total = len(S.ord_queue)

            S.num_queue = [c for c in selected if S.col_types[c] == "numeric"]
            S.num_total = len(S.num_queue)
            if S.ord_queue:
                S.cur_ord = None
                go(3)
            elif S.num_queue:
                S.cur_num = None
                go(3)
            else:
                go(4)

# ──────────────────────────────────────────────────────────────
#  STEP 3 — Configure numeric columns
#
#  FIX: Threshold group labels auto-track op + value.
#  Mechanism: store (op, thresh_val) snapshot in session state.
#  When it changes, delete the label widget keys so Streamlit
#  resets them to the new auto-computed `value=` default.
#  The user can still type a custom label — it only resets on
#  the NEXT op/value change.
#
#  FIX: Back button puts current column back into the queue
#  so navigating back to step 2 and forward again doesn't skip it.
# ──────────────────────────────────────────────────────────────
def step_configure():
    # ── Route: ordered_categorical first, then numeric ─────────
    # If there are still ordered cols to configure, handle those first.
    if S.ord_queue:
        _configure_ordered()
    elif S.num_queue:
        _configure_numeric()
    else:
        go(4)

# ──────────────────────────────────────────────────────────────
#  Step 3a — Configure ordered_categorical columns
#  User sees the auto-detected entries (e.g. "1st line", "2nd line")
#  and can rename the last one (e.g. "≥5th line") and reorder.
# ──────────────────────────────────────────────────────────────
def _configure_ordered():
    if S.cur_ord is None or S.cur_ord not in S.ord_queue:
        S.cur_ord = S.ord_queue[0]

    col  = S.cur_ord
    df   = S.df
    done = S.ord_total - len(S.ord_queue) + 1

    st.markdown(
        f'<p class="section-title">Configure: {col}</p>'
        f'<p class="section-sub">Ordered column {done} of {S.ord_total} — '
        f'set a label for each value</p>',
        unsafe_allow_html=True)

    cfg     = S.configs.get(col, {"type": "ordered_categorical", "label": col})
    entries = S.entries.get(col, build_entries(col, df, "ordered_categorical"))
    total   = len(df[col].dropna())

    # Section heading
    new_heading = st.text_input("Section heading in table:", value=cfg["label"],
                                key=f"ord_heading_{col}")
    cfg["label"] = new_heading

    st.markdown("---")
    st.markdown(
        "**Set a display label for each value.**  "
        "Reorder with ▲ ▼ — drag the last category up if needed.  \n"
        "_Tip: rename the last row to_ `≥5th line` _if it's a collapsed group._")

    for idx, entry in enumerate(entries):
        # Count per value
        cnt = (df[col].astype(str) == str(entry["key"])).sum()
        pct = cnt / total * 100 if total else 0

        r = st.columns([0.5, 1.2, 2.8, 0.45, 0.45])

        r[0].markdown(
            f"<div style='padding-top:8px;color:#aaa;font-size:12px'>{idx+1}.</div>",
            unsafe_allow_html=True)
        r[1].markdown(
            f"<div style='padding-top:8px;font-size:13px;color:#555'>"
            f"<b>{entry['key']}</b>"
            f"<span style='color:#aaa;font-size:11px;margin-left:4px'>"
            f"n={cnt}</span></div>",
            unsafe_allow_html=True)

        new_lbl = r[2].text_input(
            "Label", value=entry["label"],
            key=f"ord_lbl_{col}_{idx}",
            label_visibility="collapsed")
        entries[idx]["label"] = new_lbl

        if r[3].button("▲", key=f"ord_up_{col}_{idx}", disabled=(idx == 0)):
            entries[idx], entries[idx-1] = entries[idx-1], entries[idx]
            S.entries[col] = entries; st.rerun()
        if r[4].button("▼", key=f"ord_dn_{col}_{idx}",
                       disabled=(idx == len(entries)-1)):
            entries[idx], entries[idx+1] = entries[idx+1], entries[idx]
            S.entries[col] = entries; st.rerun()

    S.entries[col] = entries
    S.configs[col] = cfg

    st.markdown("---")
    col_l, col_r = st.columns([1, 5])
    with col_l:
        if st.button("← Back", key=f"ord_back_{col}"):
            if col not in S.ord_queue:
                S.ord_queue.insert(0, col)
            S.cur_ord = None
            go(2)
    with col_r:
        if st.button("Next →", type="primary", key=f"ord_next_{col}"):
            S.ord_queue.pop(0)
            S.cur_ord = S.ord_queue[0] if S.ord_queue else None
            if S.ord_queue:
                st.rerun()
            elif S.num_queue:
                S.cur_num = None
                go(3)   # stay on step 3 for numeric
            else:
                go(4)

# ──────────────────────────────────────────────────────────────
#  Step 3b — Configure numeric columns (unchanged logic)
# ──────────────────────────────────────────────────────────────
def _configure_numeric():
    if not S.num_queue:
        go(4); return
    if S.cur_num is None or S.cur_num not in S.num_queue:
        S.cur_num = S.num_queue[0]

    col  = S.cur_num
    df   = S.df
    s    = df[col].dropna()
    done = S.num_total - len(S.num_queue) + 1

    st.markdown(
        f'<p class="section-title">Configure: {col}</p>'
        f'<p class="section-sub">Numeric column {done} of {S.num_total}</p>',
        unsafe_allow_html=True)

    st.markdown(
        f'<div class="stats-strip">'
        f'<div class="stat-item"><div class="stat-val">{s.min():.2f}</div><div class="stat-lbl">Min</div></div>'
        f'<div class="stat-item"><div class="stat-val">{s.max():.2f}</div><div class="stat-lbl">Max</div></div>'
        f'<div class="stat-item"><div class="stat-val">{s.mean():.2f}</div><div class="stat-lbl">Mean</div></div>'
        f'<div class="stat-item"><div class="stat-val">{s.median():.2f}</div><div class="stat-lbl">Median</div></div>'
        f'</div>', unsafe_allow_html=True)

    st.markdown("**Select one or more presentation modes:**")

    use_mean   = st.checkbox("Mean ± SD",          key=f"m_mean_{col}")
    use_median = st.checkbox("Median [IQR]",        key=f"m_med_{col}")
    use_thresh = st.checkbox("Groups by threshold", key=f"m_thr_{col}")

    lbl_mean = st.text_input("Label for Mean ± SD row:",
                             value=f"{col} (Mean ± SD)",
                             key=f"lbl_mean_{col}",
                             disabled=not use_mean)
    lbl_median = st.text_input("Label for Median row:",
                               value=f"{col} (Median)",
                               key=f"lbl_med_{col}",
                               disabled=not use_median)

    lbl_thresh = col
    thresh_val = None
    op = "<"
    auto1 = auto2 = ""
    lbl1 = lbl2 = ""

    if use_thresh:
        st.markdown("---")
        st.markdown("**Threshold settings:**")

        lbl_thresh = st.text_input("Label for threshold row:",
                                   value=col, key=f"lbl_thr_{col}")

        thresh_val = st.number_input("Threshold value:",
                                     value=float(s.median()),
                                     key=f"thresh_{col}")

        op = st.radio(
            "First group condition:",
            options=["<", "<=", ">", ">="],
            format_func=lambda x: {
                "<":  "Smaller than  ( < )",
                "<=": "Smaller than or equal to  ( ≤ )",
                ">":  "Greater than  ( > )",
                ">=": "Greater than or equal to  ( ≥ )",
            }[x],
            key=f"op_{col}",
            horizontal=False,
        )

        # Auto-labels that reflect current op + threshold value
        auto1 = f"{OSYM.get(op, '?')} {fmt_thresh(thresh_val)}"
        auto2 = f"{CSYM.get(op, '?')} {fmt_thresh(thresh_val)}"

        # Detect if op or threshold changed — if so, reset label widgets
        snap_key = f"_tsnap_{col}"
        prev_snap = S.get(snap_key)
        cur_snap  = (op, thresh_val)
        if prev_snap != cur_snap:
            # Delete widget state so `value=auto` kicks in on next render
            for wk in (f"l1_{col}", f"l2_{col}"):
                if wk in st.session_state:
                    del st.session_state[wk]
            st.session_state[snap_key] = cur_snap

        lbl1 = st.text_input("Group 1 label:", value=auto1, key=f"l1_{col}")
        lbl2 = st.text_input("Group 2 label:", value=auto2, key=f"l2_{col}")

    st.markdown("---")
    col_l, col_r = st.columns([1, 5])
    with col_l:
        if st.button("← Back"):
            # Ensure current col stays at front of queue
            if col not in S.num_queue:
                S.num_queue.insert(0, col)
            S.cur_num = None
            # If there were ordered cols, go back to configure those; else step 2
            if S.ord_total > 0:
                S.ord_queue = [c for c in S.selected if S.col_types[c] == "ordered_categorical"]
                S.cur_ord = None
                go(3)
            else:
                go(2)
    with col_r:
        if st.button("Next →", type="primary", key=f"next_{col}"):
            if not (use_mean or use_median or use_thresh):
                st.error("Please select at least one presentation mode."); return
            modes = []
            if use_mean:
                modes.append({"mode": "mean_sd",
                              "label": lbl_mean.strip() or f"{col} (Mean ± SD)"})
            if use_median:
                modes.append({"mode": "median",
                              "label": lbl_median.strip() or f"{col} (Median)"})
            if use_thresh:
                if thresh_val is None:
                    st.error("Enter a threshold value."); return
                modes.append({
                    "mode":        "threshold",
                    "label":       lbl_thresh.strip() or col,
                    "op":          op,
                    "threshold":   float(thresh_val),
                    "label_true":  lbl1.strip() or auto1,
                    "label_false": lbl2.strip() or auto2,
                })
            S.configs[col] = {"type": "numeric", "label": col, "modes": modes}
            S.num_queue.pop(0)
            S.cur_num = S.num_queue[0] if S.num_queue else None
            if S.num_queue:
                st.rerun()
            else:
                go(4)

# ──────────────────────────────────────────────────────────────
#  STEP 4 — Verify labels + reorder
#
#  FIX: Both categorical and multi_value now use the unified
#  S.entries[col] list. Each entry is {"key", "label"}.
#  Reordering swaps entire dicts, so key and label always stay
#  together — no drift possible between value order and label order.
#  Both types show identical UI: raw key + editable label + ▲▼.
#
#  FIX: Back button from Verify restores num_queue so users can
#  re-configure numeric columns.
# ──────────────────────────────────────────────────────────────
def step_verify():
    df = S.df
    st.markdown('<p class="section-title">Verify & arrange</p>', unsafe_allow_html=True)
    st.markdown('<p class="section-sub">Edit labels and reorder categories before exporting.</p>',
                unsafe_allow_html=True)

    for col in S.selected:
        cfg   = S.configs[col]
        dtype = S.col_types[col]
        icon  = {"categorical": "🔤", "multi_value": "🏷️",
                 "ordered_categorical": "🔢🔤", "numeric": "🔢"}.get(dtype, "•")

        with st.expander(f"{icon}  {col}", expanded=True):

            if dtype in ("categorical", "multi_value", "ordered_categorical"):
                # Editable section heading
                new_label = st.text_input(
                    "Section heading:", value=cfg["label"], key=f"v_lbl_{col}")
                cfg["label"] = new_label

                entries = S.entries.get(col, [])
                total   = len(df[col].dropna())

                if dtype == "categorical":
                    st.markdown("**Categories** — rename and reorder with ▲ ▼:")
                    def get_cnt(key):
                        return (df[col].astype(str) == str(key)).sum()
                elif dtype == "ordered_categorical":
                    st.markdown("**Values** — rename (e.g. last row → `≥5th line`) and reorder with ▲ ▼:")
                    def get_cnt(key):
                        return (df[col].astype(str) == str(key)).sum()
                else:
                    counts = multi_value_counts(df[col])
                    st.markdown(
                        "**Terms** — rename and reorder with ▲ ▼.  "
                        "Percentages are per patient (can sum >100%).")
                    def get_cnt(key):
                        return counts.get(key, 0)

                for idx, entry in enumerate(entries):
                    cnt = get_cnt(entry["key"])
                    pct = cnt / total * 100 if total else 0

                    r = st.columns([0.5, 2.8, 2.8, 0.45, 0.45])

                    r[0].markdown(
                        f"<div style='padding-top:8px;color:#aaa;font-size:12px'>"
                        f"{idx + 1}.</div>",
                        unsafe_allow_html=True)

                    r[1].markdown(
                        f"<div style='padding-top:8px;font-size:13px;color:#555'>"
                        f"{entry['key']}"
                        f"<span style='color:#aaa;font-size:11px;margin-left:6px'>"
                        f"n={cnt} · {fmt_pct(pct)}</span></div>",
                        unsafe_allow_html=True)

                    new_lbl = r[2].text_input(
                        "Label", value=entry["label"],
                        key=f"entry_lbl_{col}_{idx}",
                        label_visibility="collapsed")
                    entries[idx]["label"] = new_lbl   # write back immediately

                    if r[3].button("▲", key=f"up_{col}_{idx}", disabled=(idx == 0)):
                        entries[idx], entries[idx - 1] = entries[idx - 1], entries[idx]
                        S.entries[col] = entries
                        st.rerun()

                    if r[4].button("▼", key=f"dn_{col}_{idx}",
                                   disabled=(idx == len(entries) - 1)):
                        entries[idx], entries[idx + 1] = entries[idx + 1], entries[idx]
                        S.entries[col] = entries
                        st.rerun()

                S.entries[col] = entries

            else:  # numeric
                mode_names = {
                    "mean_sd":   "Mean ± SD",
                    "median":    "Median [IQR]",
                    "threshold": "Threshold groups",
                }
                for idx, mc in enumerate(cfg["modes"]):
                    new_lbl = st.text_input(
                        f"Label — {mode_names.get(mc['mode'], mc['mode'])}:",
                        value=mc["label"],
                        key=f"v_num_{col}_{idx}")
                    mc["label"] = new_lbl
                    if mc["mode"] == "threshold":
                        st.caption(
                            f"Groups: **{mc['label_true']}** "
                            f"({OSYM.get(mc['op'], '')} {fmt_thresh(mc['threshold'])})"
                            f"  ·  **{mc['label_false']}** "
                            f"({CSYM.get(mc['op'], '')} {fmt_thresh(mc['threshold'])})")

    st.markdown("---")
    col_l, col_r = st.columns([1, 5])
    with col_l:
        if st.button("← Back"):
            has_numeric = any(S.col_types[c] == "numeric" for c in S.selected)
            has_ordered = any(S.col_types[c] == "ordered_categorical" for c in S.selected)
            if has_numeric:
                S.num_queue = [c for c in S.selected if S.col_types[c] == "numeric"]
                S.cur_num   = None
                go(3)
            elif has_ordered:
                S.ord_queue = [c for c in S.selected if S.col_types[c] == "ordered_categorical"]
                S.cur_ord   = None
                go(3)
            else:
                go(2)
    with col_r:
        if st.button("Build Word table →", type="primary"):
            S.word_bytes = build_word_bytes(S.selected, S.configs, S.entries, df)
            go(5)

# ──────────────────────────────────────────────────────────────
#  STEP 5 — Download
# ──────────────────────────────────────────────────────────────
def step_download():
    if not S.word_bytes:
        S.word_bytes = build_word_bytes(S.selected, S.configs, S.entries, S.df)

    st.markdown(
        '<div class="success-box">'
        '<div class="success-check">✓</div>'
        '<div class="success-title">Your table is ready</div>'
        '<div class="success-sub">Click the button below to download the Word file.</div>'
        '</div>',
        unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    st.download_button(
        label="📄  Download Word file",
        data=S.word_bytes,
        file_name="summary_table.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

    st.markdown("<br>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        if st.button("← Edit columns", use_container_width=True):
            S.word_bytes = None
            go(4)
    with col2:
        if st.button("🔄  Start over", use_container_width=True):
            for key in list(st.session_state.keys()):
                del st.session_state[key]
            st.rerun()

# ──────────────────────────────────────────────────────────────
#  Router
# ──────────────────────────────────────────────────────────────
render_steps()

if   S.step == 1: step_upload()
elif S.step == 2: step_select()
elif S.step == 3: step_configure()
elif S.step == 4: step_verify()
elif S.step == 5: step_download()
